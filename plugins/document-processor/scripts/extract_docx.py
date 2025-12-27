#!/usr/bin/env python3
"""
Extract text and convert DOCX files to markdown.

Usage:
    python extract_docx.py --file document.docx
    python extract_docx.py --file document.docx --markdown
    python extract_docx.py --file document.docx --extract-images --output-dir ./extracted/
    python extract_docx.py --file document.docx --metadata-only
    python extract_docx.py --file document.docx --json

Features:
- Full text extraction preserving structure
- Markdown conversion (headings, lists, tables, bold/italic)
- Image extraction to separate directory
- Table extraction as markdown tables
- Metadata extraction (author, created, modified)
- JSON output mode for programmatic use
"""

import argparse
import base64
import json
import os
import re
import sys
from dataclasses import dataclass, asdict, field
from datetime import datetime
from pathlib import Path
from typing import Optional, List
from zipfile import ZipFile


@dataclass
class DocxMetadata:
    """DOCX document metadata."""
    title: Optional[str] = None
    author: Optional[str] = None
    subject: Optional[str] = None
    keywords: Optional[str] = None
    created: Optional[str] = None
    modified: Optional[str] = None
    last_modified_by: Optional[str] = None
    revision: Optional[int] = None
    paragraph_count: int = 0
    word_count: int = 0
    char_count: int = 0
    table_count: int = 0
    image_count: int = 0
    file_size_bytes: int = 0
    file_size_human: str = ""


@dataclass
class ExtractedImage:
    """Information about an extracted image."""
    filename: str
    original_name: str
    content_type: str
    size_bytes: int
    saved_path: Optional[str] = None


@dataclass
class ExtractionResult:
    """Complete extraction result."""
    success: bool
    file_path: str
    metadata: DocxMetadata
    content: str
    images: List[ExtractedImage] = field(default_factory=list)
    error: Optional[str] = None


def format_file_size(size_bytes: int) -> str:
    """Format bytes to human-readable size."""
    for unit in ['B', 'KB', 'MB', 'GB']:
        if size_bytes < 1024:
            return f"{size_bytes:.1f} {unit}"
        size_bytes /= 1024
    return f"{size_bytes:.1f} TB"


def format_datetime(dt) -> Optional[str]:
    """Format datetime to ISO string."""
    if dt is None:
        return None
    if isinstance(dt, datetime):
        return dt.isoformat()
    return str(dt)


def extract_metadata(doc) -> DocxMetadata:
    """Extract metadata from DOCX document."""
    core_props = doc.core_properties
    file_stat = os.stat(doc._part.package.name if hasattr(doc._part.package, 'name') else "")

    # Count elements
    para_count = len(doc.paragraphs)
    table_count = len(doc.tables)

    # Count words and characters
    all_text = "\n".join(p.text for p in doc.paragraphs)
    word_count = len(all_text.split())
    char_count = len(all_text)

    # Count images
    image_count = 0
    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            image_count += 1

    return DocxMetadata(
        title=core_props.title,
        author=core_props.author,
        subject=core_props.subject,
        keywords=core_props.keywords,
        created=format_datetime(core_props.created),
        modified=format_datetime(core_props.modified),
        last_modified_by=core_props.last_modified_by,
        revision=core_props.revision,
        paragraph_count=para_count,
        word_count=word_count,
        char_count=char_count,
        table_count=table_count,
        image_count=image_count,
        file_size_bytes=file_stat.st_size if file_stat else 0,
        file_size_human=format_file_size(file_stat.st_size) if file_stat else ""
    )


def get_paragraph_style(paragraph) -> str:
    """Get the style name of a paragraph."""
    if paragraph.style and paragraph.style.name:
        return paragraph.style.name
    return "Normal"


def run_to_markdown(run) -> str:
    """Convert a run to markdown with formatting."""
    text = run.text
    if not text:
        return ""

    # Apply formatting
    if run.bold:
        text = f"**{text}**"
    if run.italic:
        text = f"*{text}*"
    if run.underline:
        text = f"<u>{text}</u>"
    if run.font.strike:
        text = f"~~{text}~~"

    return text


def paragraph_to_markdown(paragraph) -> str:
    """Convert a paragraph to markdown."""
    style = get_paragraph_style(paragraph)
    text_parts = [run_to_markdown(run) for run in paragraph.runs]
    text = "".join(text_parts).strip()

    if not text:
        return ""

    # Handle heading styles
    if style.startswith("Heading"):
        try:
            level = int(style.replace("Heading", "").strip())
            level = min(level, 6)  # Markdown supports h1-h6
            return "#" * level + " " + text
        except ValueError:
            pass

    # Handle list styles
    if style.startswith("List"):
        # Check if it's a numbered list
        if "Number" in style or "Bullet" not in style:
            return f"1. {text}"
        return f"- {text}"

    # Handle quote style
    if "Quote" in style:
        return f"> {text}"

    return text


def table_to_markdown(table) -> str:
    """Convert a table to markdown format."""
    if not table.rows:
        return ""

    # Extract all cell contents
    rows_data = []
    max_cols = 0
    for row in table.rows:
        row_data = []
        for cell in row.cells:
            cell_text = " ".join(p.text.strip() for p in cell.paragraphs)
            cell_text = cell_text.replace("|", "\\|")  # Escape pipes
            row_data.append(cell_text)
        rows_data.append(row_data)
        max_cols = max(max_cols, len(row_data))

    if not rows_data:
        return ""

    # Normalize column count
    for row in rows_data:
        while len(row) < max_cols:
            row.append("")

    # Build markdown table
    lines = []

    # Header row
    header = "| " + " | ".join(rows_data[0]) + " |"
    lines.append(header)

    # Separator row
    separator = "| " + " | ".join(["---"] * max_cols) + " |"
    lines.append(separator)

    # Data rows
    for row in rows_data[1:]:
        line = "| " + " | ".join(row) + " |"
        lines.append(line)

    return "\n".join(lines)


def extract_images(doc, output_dir: str) -> List[ExtractedImage]:
    """Extract all images from the document."""
    images = []
    os.makedirs(output_dir, exist_ok=True)

    for rel_id, rel in doc.part.rels.items():
        if "image" in rel.target_ref:
            try:
                image_part = rel.target_part
                image_bytes = image_part.blob
                original_name = os.path.basename(rel.target_ref)

                # Determine extension from content type
                content_type = image_part.content_type
                ext_map = {
                    "image/png": ".png",
                    "image/jpeg": ".jpg",
                    "image/gif": ".gif",
                    "image/bmp": ".bmp",
                    "image/tiff": ".tiff",
                    "image/x-emf": ".emf",
                    "image/x-wmf": ".wmf"
                }
                ext = ext_map.get(content_type, ".bin")

                # Create unique filename
                filename = f"image_{len(images) + 1}{ext}"
                save_path = os.path.join(output_dir, filename)

                # Save image
                with open(save_path, "wb") as f:
                    f.write(image_bytes)

                images.append(ExtractedImage(
                    filename=filename,
                    original_name=original_name,
                    content_type=content_type,
                    size_bytes=len(image_bytes),
                    saved_path=save_path
                ))

            except Exception as e:
                print(f"Warning: Could not extract image {rel.target_ref}: {e}", file=sys.stderr)

    return images


def extract_text_plain(doc) -> str:
    """Extract plain text from document."""
    lines = []

    for element in doc.element.body:
        # Handle paragraphs
        if element.tag.endswith('p'):
            for para in doc.paragraphs:
                if para._element == element:
                    text = para.text.strip()
                    if text:
                        lines.append(text)
                    break

        # Handle tables
        elif element.tag.endswith('tbl'):
            for table in doc.tables:
                if table._element == element:
                    for row in table.rows:
                        row_text = " | ".join(cell.text.strip() for cell in row.cells)
                        if row_text.replace("|", "").strip():
                            lines.append(row_text)
                    lines.append("")  # Blank line after table
                    break

    return "\n".join(lines)


def extract_markdown(doc, image_dir: Optional[str] = None) -> str:
    """Extract document content as markdown."""
    lines = []
    image_counter = 0

    # Process document body in order
    for element in doc.element.body:
        # Handle paragraphs
        if element.tag.endswith('p'):
            for para in doc.paragraphs:
                if para._element == element:
                    md = paragraph_to_markdown(para)
                    if md:
                        lines.append(md)
                    else:
                        lines.append("")  # Preserve empty paragraphs as spacing
                    break

        # Handle tables
        elif element.tag.endswith('tbl'):
            for table in doc.tables:
                if table._element == element:
                    md_table = table_to_markdown(table)
                    if md_table:
                        lines.append("")
                        lines.append(md_table)
                        lines.append("")
                    break

    # Clean up multiple blank lines
    content = "\n".join(lines)
    content = re.sub(r'\n{3,}', '\n\n', content)

    return content.strip()


def extract_docx(file_path: str, as_markdown: bool = False,
                 extract_images_to: Optional[str] = None,
                 metadata_only: bool = False) -> ExtractionResult:
    """Extract content from DOCX file."""
    try:
        from docx import Document
    except ImportError:
        return ExtractionResult(
            success=False,
            file_path=file_path,
            metadata=DocxMetadata(),
            content="",
            error="python-docx not installed. Run: pip install python-docx"
        )

    try:
        doc = Document(file_path)

        # Store file path for metadata extraction
        doc._part.package.name = file_path

        # Extract metadata
        metadata = extract_metadata(doc)

        if metadata_only:
            return ExtractionResult(
                success=True,
                file_path=file_path,
                metadata=metadata,
                content=""
            )

        # Extract images if requested
        images = []
        if extract_images_to:
            images = extract_images(doc, extract_images_to)
            metadata.image_count = len(images)

        # Extract content
        if as_markdown:
            content = extract_markdown(doc, extract_images_to)
        else:
            content = extract_text_plain(doc)

        # Update metadata with actual counts
        metadata.word_count = len(content.split())
        metadata.char_count = len(content)

        return ExtractionResult(
            success=True,
            file_path=file_path,
            metadata=metadata,
            content=content,
            images=images
        )

    except Exception as e:
        return ExtractionResult(
            success=False,
            file_path=file_path,
            metadata=DocxMetadata(),
            content="",
            error=str(e)
        )


def format_human_output(result: ExtractionResult) -> str:
    """Format extraction result for human reading."""
    lines = [
        "=" * 60,
        "DOCX EXTRACTION REPORT",
        "=" * 60,
        "",
        f"File: {result.file_path}",
        f"Status: {'SUCCESS' if result.success else 'FAILED'}",
    ]

    if result.error:
        lines.append(f"Error: {result.error}")
        return "\n".join(lines)

    # Metadata section
    lines.extend([
        "",
        "METADATA",
        "-" * 40
    ])

    meta = result.metadata
    if meta.title:
        lines.append(f"Title: {meta.title}")
    if meta.author:
        lines.append(f"Author: {meta.author}")
    if meta.subject:
        lines.append(f"Subject: {meta.subject}")
    if meta.created:
        lines.append(f"Created: {meta.created}")
    if meta.modified:
        lines.append(f"Modified: {meta.modified}")
    if meta.last_modified_by:
        lines.append(f"Last Modified By: {meta.last_modified_by}")

    lines.extend([
        f"Paragraphs: {meta.paragraph_count}",
        f"Tables: {meta.table_count}",
        f"Images: {meta.image_count}",
        f"Words: {meta.word_count:,}",
        f"Characters: {meta.char_count:,}",
        f"Size: {meta.file_size_human}"
    ])

    # Images section
    if result.images:
        lines.extend([
            "",
            "EXTRACTED IMAGES",
            "-" * 40
        ])
        for img in result.images:
            lines.append(f"  {img.filename} ({format_file_size(img.size_bytes)}) -> {img.saved_path}")

    # Content section
    if result.content:
        lines.extend([
            "",
            "CONTENT",
            "-" * 40,
            result.content
        ])

    lines.append("")
    lines.append("=" * 60)
    return "\n".join(lines)


def format_json_output(result: ExtractionResult) -> str:
    """Format extraction result as JSON."""
    output = {
        "success": result.success,
        "file_path": result.file_path,
        "metadata": asdict(result.metadata)
    }

    if result.error:
        output["error"] = result.error
    else:
        output["content"] = result.content
        if result.images:
            output["images"] = [asdict(img) for img in result.images]

    return json.dumps(output, indent=2, ensure_ascii=False)


def main():
    parser = argparse.ArgumentParser(
        description="Extract text and convert DOCX files to markdown",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  python extract_docx.py --file document.docx
  python extract_docx.py --file document.docx --markdown
  python extract_docx.py --file document.docx --extract-images --output-dir ./images/
  python extract_docx.py --file document.docx --metadata-only
  python extract_docx.py --file document.docx --json
        """
    )

    parser.add_argument("--file", "-f", required=True, help="DOCX file to extract")
    parser.add_argument("--markdown", "-md", action="store_true", help="Convert to markdown format")
    parser.add_argument("--extract-images", "-i", action="store_true", help="Extract embedded images")
    parser.add_argument("--output-dir", "-o", default="./extracted_images", help="Directory for extracted images")
    parser.add_argument("--metadata-only", "-m", action="store_true", help="Extract only metadata")
    parser.add_argument("--json", "-j", action="store_true", help="Output as JSON")

    args = parser.parse_args()

    # Validate file exists
    if not os.path.exists(args.file):
        print(f"Error: File not found: {args.file}", file=sys.stderr)
        sys.exit(1)

    # Validate file extension
    if not args.file.lower().endswith('.docx'):
        print(f"Warning: File does not have .docx extension", file=sys.stderr)

    # Extract content
    result = extract_docx(
        file_path=args.file,
        as_markdown=args.markdown,
        extract_images_to=args.output_dir if args.extract_images else None,
        metadata_only=args.metadata_only
    )

    # Format and output
    if args.json:
        print(format_json_output(result))
    else:
        print(format_human_output(result))

    # Exit code
    sys.exit(0 if result.success else 1)


if __name__ == "__main__":
    main()
